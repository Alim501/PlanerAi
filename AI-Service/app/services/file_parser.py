"""
File parsing service for extracting text from various file formats
"""
import logging
from pathlib import Path
from typing import Optional
import PyPDF2
from docx import Document
from PIL import Image
import pytesseract

logger = logging.getLogger(__name__)


class FileParserService:
    """Service for parsing different file types"""

    async def parse_file(self, file_path: str, file_type: str) -> str:
        """
        Parse file and extract text content

        Args:
            file_path: Path to the file
            file_type: Type of file (pdf, docx, image)

        Returns:
            Extracted text content
        """
        try:
            file_type = file_type.lower()

            if file_type == "pdf":
                return await self._parse_pdf(file_path)
            elif file_type == "docx":
                return await self._parse_docx(file_path)
            elif file_type in ["image", "jpg", "jpeg", "png"]:
                return await self._parse_image(file_path)
            elif file_type == "txt":
                return await self._parse_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {e}")
            raise

    async def _parse_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text_content = []

            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)

                logger.info(f"Parsing PDF with {num_pages} pages")

                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)

            result = "\n\n".join(text_content)
            logger.info(f"Extracted {len(result)} characters from PDF")
            return result

        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            raise

    async def _parse_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_content = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)

            result = "\n\n".join(text_content)
            logger.info(f"Extracted {len(result)} characters from DOCX")
            return result

        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise

    async def _parse_image(self, file_path: str) -> str:
        """Extract text from image using OCR"""
        try:
            image = Image.open(file_path)

            # Perform OCR
            # Note: Requires tesseract to be installed on the system
            # Install: brew install tesseract (macOS) or apt-get install tesseract-ocr (Linux)
            text = pytesseract.image_to_string(image, lang="rus+eng")

            logger.info(f"Extracted {len(text)} characters from image using OCR")
            return text.strip()

        except Exception as e:
            logger.error(f"Error parsing image: {e}")
            # If tesseract is not installed, return a warning message
            if "tesseract" in str(e).lower():
                return "[OCR not available - tesseract not installed]"
            raise

    async def _parse_txt(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                result = file.read()
            logger.info(f"Extracted {len(result)} characters from TXT")
            return result
        except Exception as e:
            logger.error(f"Error parsing TXT: {e}")
            raise

    def get_file_info(self, file_path: str) -> dict:
        """Get basic file information"""
        path = Path(file_path)

        return {
            "filename": path.name,
            "extension": path.suffix,
            "size_bytes": path.stat().st_size if path.exists() else 0,
            "exists": path.exists(),
        }


# Singleton instance
_file_parser: Optional[FileParserService] = None


def get_file_parser() -> FileParserService:
    """Get singleton file parser instance"""
    global _file_parser
    if _file_parser is None:
        _file_parser = FileParserService()
    return _file_parser
